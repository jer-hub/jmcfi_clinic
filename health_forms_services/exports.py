"""
Document export helpers — generates .docx files that replicate the official
JMCFI Clinic PDF form layouts, populated with data from the database.

Each ``generate_*`` function accepts a model instance and returns a
``python-docx`` ``Document`` ready to be saved or streamed as a response.
"""

import os
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Path to the folder that contains the official .docx templates
TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "health_forms"


# ─── Shared constants ─────────────────────────────────────────────

INSTITUTION_NAME = "JOSE MARIA COLLEGE"
INSTITUTION_SUB = "FOUNDATION, INC."
INSTITUTION_ADDRESS = "Philippine-Japan Friendship Highway, Sasa, Davao City, 8000"
DEPARTMENT = "HEALTH SERVICES"
REVISION = "Rev: 00.01.Feb.2020"

FONT_NAME = "Arial"
DARK_FILL = "1a1a1a"
LIGHT_FILL = "E5E5E5"


# ─── Low-level helpers ────────────────────────────────────────────

def _val(value, default=""):
    """Return a printable string; handles None, empty, booleans, dates."""
    if value is None or value == "":
        return str(default)
    if isinstance(value, bool):
        return "Yes" if value else "No"
    return str(value)


def _set_cell_shading(cell, color_hex):
    """Apply background colour to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Set borders on a cell.  kwargs: top, bottom, left, right
    each value is a dict with keys: sz, val, color  (e.g. sz='4', val='single', color='000000')
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(element)


def _set_table_borders(table):
    """Apply thin black borders to every cell in the table."""
    border = {"sz": "4", "val": "single", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top=border, bottom=border, left=border, right=border)


def _set_narrow_row_height(row, height_twips=300):
    """Make a row height fixed/tight (in twips — 1440 = 1 inch)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{height_twips}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def _run(paragraph, text, bold=False, size=Pt(9), color=None, font_name=FONT_NAME):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def _cell_text(cell, text, bold=False, size=Pt(9), alignment=None, color=None):
    """Write styled text into a table cell."""
    p = cell.paragraphs[0]
    p.clear()
    if alignment is not None:
        p.alignment = alignment
    _run(p, str(text), bold=bold, size=size, color=color)


def _label_value_cell(cell, label, value, label_size=Pt(7), value_size=Pt(9)):
    """Small uppercase label + value below it in one cell."""
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, label.upper() + "\n", bold=True, size=label_size, color=RGBColor(0x66, 0x66, 0x66))
    _run(p, _val(value), bold=False, size=value_size)


# ─── Letterhead ───────────────────────────────────────────────────

def _add_letterhead(doc, form_title, form_number=None):
    """Add the official JMCFI letterhead centred at the top of the document."""
    for text, size, bold in [
        (INSTITUTION_NAME, Pt(14), True),
        (INSTITUTION_SUB, Pt(12), True),
        (INSTITUTION_ADDRESS, Pt(8), False),
        (DEPARTMENT, Pt(10), True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, text, bold=bold, size=size)

    # Form title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    _run(p, form_title.upper(), bold=True, size=Pt(13))

    # Thin rule under header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_footer(doc, form_number):
    """Add form number and revision as footer text at the bottom."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)

    tab_stops = p.paragraph_format.tab_stops
    from docx.shared import Inches
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    _run(p, form_number, size=Pt(7.5), color=RGBColor(0x66, 0x66, 0x66))
    _run(p, "\t", size=Pt(7.5))
    _run(p, REVISION, size=Pt(7.5), color=RGBColor(0x66, 0x66, 0x66))


# ─── Section header (dark bar) ───────────────────────────────────

def _add_section_header(doc, title):
    """Dark bar section divider matching PDF style."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, DARK_FILL)
    _cell_text(cell, title.upper(), bold=True, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))
    _set_table_borders(table)

    # tiny spacing after
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)


# ─── Data grid (label/value pairs as bordered table) ─────────────

def _add_data_grid(doc, fields, cols=3):
    """Render a list of (label, value) pairs as a bordered grid table.
    ``fields`` is a list of (label, value) tuples.
    """
    rows_needed = -(-len(fields) // cols)  # ceil division
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, value) in enumerate(fields):
        r, c = divmod(idx, cols)
        cell = table.cell(r, c)
        _label_value_cell(cell, label, value)
    _set_table_borders(table)

    # spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)


# ─── Signature block ─────────────────────────────────────────────

def _add_signature_block(doc, name="", title_text="Physician", license_label="License No.", license_no="", ptr_label="PTR No.", ptr_no=""):
    """Blank signature line with name/license underneath, right-aligned."""
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Signature line — right aligned
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    _run(p, "_" * 40, size=Pt(9))

    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        _run(p, _val(name), bold=True, size=Pt(10))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _run(p, title_text, size=Pt(8), color=RGBColor(0x66, 0x66, 0x66))

    if license_no:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, f"{license_label}: {license_no}", size=Pt(8), color=RGBColor(0x66, 0x66, 0x66))

    if ptr_no:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _run(p, f"{ptr_label}: {ptr_no}", size=Pt(8), color=RGBColor(0x66, 0x66, 0x66))


# ─── Page setup helper ────────────────────────────────────────────

def _setup_page(doc, orient="portrait"):
    """Configure A4 page with narrow margins."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    if orient == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


# ═══════════════════════════════════════════════════════════════════
#  Form-specific generators
# ═══════════════════════════════════════════════════════════════════


def generate_health_profile(form):
    """F-HSS-20-0001 — Health Profile Form."""
    doc = Document()
    _setup_page(doc)
    _add_letterhead(doc, "Health Profile Form", "F-HSS-20-0001")

    # ── Personal Information ──
    _add_section_header(doc, "Personal Information")
    _add_data_grid(doc, [
        ("Last Name", form.last_name),
        ("First Name", form.first_name),
        ("Middle Name", form.middle_name),
        ("Permanent Address", form.permanent_address),
        ("Zip Code", form.zip_code),
        ("Current Address", form.current_address),
        ("Religion", form.religion),
        ("Civil Status", form.get_civil_status_display() if form.civil_status else ""),
        ("Place of Birth", form.place_of_birth),
        ("Date of Birth", form.date_of_birth),
        ("Citizenship", form.citizenship),
        ("Age", form.age),
        ("Gender", form.get_gender_display() if form.gender else ""),
        ("Email", form.email_address),
        ("Mobile No.", form.mobile_number),
        ("Telephone No.", form.telephone_number),
        ("Designation", form.get_designation_display() if form.designation else ""),
        ("Dept / College / Office", form.department_college_office),
    ], cols=3)

    # Emergency contact
    _add_data_grid(doc, [
        ("Guardian / Emergency Contact", form.guardian_name),
        ("Contact Number", form.guardian_contact),
    ], cols=2)

    # ── Immunization Records ──
    _add_section_header(doc, "Immunization Records")
    immunizations = [
        ("COVID-19", form.immunization_covid19, form.immunization_covid19_date),
        ("Influenza", form.immunization_influenza, form.immunization_influenza_date),
        ("Pneumonia", form.immunization_pneumonia, form.immunization_pneumonia_date),
        ("Polio", form.immunization_polio, form.immunization_polio_date),
        ("Hepatitis B", form.immunization_hepatitis_b, form.immunization_hepatitis_b_date),
        ("BCG", form.immunization_bcg, form.immunization_bcg_date),
        ("DPT / Tetanus", form.immunization_dpt_tetanus, form.immunization_dpt_tetanus_date),
        ("Rotavirus", form.immunization_rotavirus, form.immunization_rotavirus_date),
        ("HIB", form.immunization_hib, form.immunization_hib_date),
        ("Measles / MMR", form.immunization_measles_mmr, form.immunization_measles_mmr_date),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(["Vaccine", "Received", "Date"]):
        cell = table.cell(0, i)
        _set_cell_shading(cell, LIGHT_FILL)
        _cell_text(cell, hdr, bold=True, size=Pt(8))
    for name, received, date in immunizations:
        row = table.add_row()
        _cell_text(row.cells[0], name, size=Pt(8))
        _cell_text(row.cells[1], "✓" if received else "—", size=Pt(9),
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row.cells[2], _val(date, "—"), size=Pt(8))
    _set_table_borders(table)

    if form.immunization_others:
        p = doc.add_paragraph()
        _run(p, "Other: ", bold=True, size=Pt(8))
        _run(p, _val(form.immunization_others), size=Pt(8))

    # spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Illnesses / Medical Conditions ──
    _add_section_header(doc, "Illnesses / Medical Conditions")
    illnesses = [
        ("Measles", form.illness_measles),
        ("Mumps", form.illness_mumps),
        ("Rubella", form.illness_rubella),
        ("Chickenpox", form.illness_chickenpox),
        ("PTB / PKI", form.illness_ptb_pki),
        ("Hypertension", form.illness_hypertension),
        ("Diabetes", form.illness_diabetes),
        ("Asthma", form.illness_asthma),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(["Condition", "Present", "Condition", "Present"]):
        _set_cell_shading(table.cell(0, i), LIGHT_FILL)
        _cell_text(table.cell(0, i), hdr, bold=True, size=Pt(8))
    for idx in range(0, len(illnesses), 2):
        row = table.add_row()
        name1, val1 = illnesses[idx]
        _cell_text(row.cells[0], name1, size=Pt(8))
        _cell_text(row.cells[1], "✓" if val1 else "—", size=Pt(9),
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if idx + 1 < len(illnesses):
            name2, val2 = illnesses[idx + 1]
            _cell_text(row.cells[2], name2, size=Pt(8))
            _cell_text(row.cells[3], "✓" if val2 else "—", size=Pt(9),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_table_borders(table)

    if form.illness_others:
        p = doc.add_paragraph()
        _run(p, "Other: ", bold=True, size=Pt(8))
        _run(p, _val(form.illness_others), size=Pt(8))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Allergies & Medications
    _add_data_grid(doc, [
        ("Allergies", form.allergies),
        ("Current Medications", form.current_medications),
    ], cols=2)

    # ── OB-GYN History ──
    if form.gender == 'female' or any([
        form.menarche_age, form.menstrual_duration, form.menstrual_interval,
        form.menstrual_amount, form.menstrual_symptoms, form.obstetric_history
    ]):
        _add_section_header(doc, "OB-GYN History")
        _add_data_grid(doc, [
            ("Menarche Age", form.menarche_age),
            ("Duration", form.menstrual_duration),
            ("Interval", form.menstrual_interval),
            ("Amount", form.menstrual_amount),
            ("Symptoms", form.menstrual_symptoms),
            ("Obstetric History", form.obstetric_history),
        ], cols=3)

    # ── Present Illness ──
    if form.present_illness:
        _add_section_header(doc, "Present Illness")
        p = doc.add_paragraph()
        _run(p, _val(form.present_illness), size=Pt(9))

    # ── Physical Examination ──
    _add_section_header(doc, "Physical Examination — Vital Signs")
    _add_data_grid(doc, [
        ("Blood Pressure", form.blood_pressure),
        ("Heart Rate", f"{form.heart_rate} bpm" if form.heart_rate else ""),
        ("Respiratory Rate", f"{form.respiratory_rate} /min" if form.respiratory_rate else ""),
        ("Temperature", f"{form.temperature} °C" if form.temperature else ""),
        ("SpO₂", f"{form.spo2}%" if form.spo2 else ""),
        ("Height", f"{form.height} m" if form.height else ""),
        ("Weight", f"{form.weight} kg" if form.weight else ""),
        ("BMI", f"{form.bmi} ({form.bmi_remarks})" if form.bmi else ""),
    ], cols=4)

    _add_section_header(doc, "Physical Examination — Findings")
    exam_fields = [
        ("General", form.exam_general),
        ("HEENT", form.exam_heent),
        ("Chest / Lungs", form.exam_chest_lungs),
        ("Abdomen", form.exam_abdomen),
        ("Genitourinary", form.exam_genitourinary),
        ("Extremities", form.exam_extremities),
        ("Neurologic", form.exam_neurologic),
        ("Other Findings", form.exam_other_findings),
    ]
    _add_data_grid(doc, exam_fields, cols=2)

    # ── Diagnostic Tests ──
    _add_section_header(doc, "Diagnostic Tests")
    tests = [
        ("Chest X-Ray", form.test_chest_xray, form.test_chest_xray_findings, form.test_chest_xray_date),
        ("CBC", form.test_cbc, form.test_cbc_findings, form.test_cbc_date),
        ("Urinalysis", form.test_urinalysis, form.test_urinalysis_findings, form.test_urinalysis_date),
        ("Drug Test", form.test_drug_test, form.test_drug_test_findings, form.test_drug_test_date),
        ("Psychological", form.test_psychological, form.test_psychological_findings, form.test_psychological_date),
        ("HBsAg", form.test_hbsag, form.test_hbsag_findings, form.test_hbsag_date),
        ("Anti-HBs Titer", form.test_anti_hbs_titer, form.test_anti_hbs_titer_findings, form.test_anti_hbs_titer_date),
        ("Fecalysis", form.test_fecalysis, form.test_fecalysis_findings, form.test_fecalysis_date),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(["Test", "Done", "Findings", "Date"]):
        _set_cell_shading(table.cell(0, i), LIGHT_FILL)
        _cell_text(table.cell(0, i), hdr, bold=True, size=Pt(8))
    for name, done, findings, date in tests:
        row = table.add_row()
        _cell_text(row.cells[0], name, size=Pt(8))
        _cell_text(row.cells[1], "✓" if done else "—", size=Pt(9),
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row.cells[2], _val(findings, "—"), size=Pt(8))
        _cell_text(row.cells[3], _val(date, "—"), size=Pt(8))
    _set_table_borders(table)

    if form.test_others:
        p = doc.add_paragraph()
        _run(p, "Other: ", bold=True, size=Pt(8))
        _run(p, _val(form.test_others), size=Pt(8))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Clinical Summary ──
    _add_section_header(doc, "Clinical Summary")
    _add_data_grid(doc, [
        ("Physician Impression", form.physician_impression),
        ("Final Remarks", form.final_remarks),
        ("Recommendations", form.recommendations),
        ("Examination Date", form.examination_date),
    ], cols=2)

    _add_signature_block(
        doc,
        name=_val(form.examining_physician, ""),
        title_text="Examining Physician",
        license_label="License No.",
        license_no="",
    )

    _add_footer(doc, "F-HSS-20-0001")
    return doc


# ──────────────────────────────────────────────────────────────────
#  F-HSS-20-0002 — Patient Chart
# ──────────────────────────────────────────────────────────────────

def generate_patient_chart(chart):
    """F-HSS-20-0002 — Patient Chart."""
    doc = Document()
    _setup_page(doc)
    _add_letterhead(doc, "Patient Chart", "F-HSS-20-0002")

    _add_section_header(doc, "Personal Information")
    _add_data_grid(doc, [
        ("Last Name", chart.last_name),
        ("First Name", chart.first_name),
        ("Middle Name", chart.middle_name),
        ("Address", chart.address),
        ("Date of Birth", chart.date_of_birth),
        ("Place of Birth", chart.place_of_birth),
        ("Age", chart.age),
        ("Gender", chart.get_gender_display() if chart.gender else ""),
        ("Civil Status", chart.get_civil_status_display() if chart.civil_status else ""),
        ("Email", chart.email_address),
        ("Contact No.", chart.contact_number),
        ("Telephone No.", chart.telephone_number),
        ("Designation", chart.get_designation_display() if chart.designation else ""),
        ("Dept / College / Office", chart.department_college_office),
    ], cols=3)

    _add_data_grid(doc, [
        ("Guardian / Emergency Contact", chart.guardian_name),
        ("Contact Number", chart.guardian_contact),
    ], cols=2)

    # ── Consultation Log ──
    _add_section_header(doc, "Consultation Log")
    entries = chart.entries.all().order_by('date_and_time')
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(["Date and Time", "Findings", "Doctor's Orders"]):
        _set_cell_shading(table.cell(0, i), LIGHT_FILL)
        _cell_text(table.cell(0, i), hdr, bold=True, size=Pt(8))

    if entries.exists():
        for entry in entries:
            row = table.add_row()
            _cell_text(row.cells[0], entry.date_and_time.strftime("%Y-%m-%d %H:%M"), size=Pt(8))
            _cell_text(row.cells[1], _val(entry.findings), size=Pt(8))
            _cell_text(row.cells[2], _val(entry.doctors_orders), size=Pt(8))
    else:
        # blank rows for manual fill
        for _ in range(8):
            row = table.add_row()
            for cell in row.cells:
                _cell_text(cell, "", size=Pt(8))
    _set_table_borders(table)

    _add_footer(doc, "F-HSS-20-0002")
    return doc


# ──────────────────────────────────────────────────────────────────
#  F-HSS-20-0003 — Dental Records
# ──────────────────────────────────────────────────────────────────

def generate_dental_form(form):
    """F-HSS-20-0003 — Dental Records Form."""
    doc = Document()
    _setup_page(doc)
    _add_letterhead(doc, "Dental Records", "F-HSS-20-0003")

    # Personal info
    _add_section_header(doc, "Personal Information")
    _add_data_grid(doc, [
        ("Last Name", form.last_name),
        ("First Name", form.first_name),
        ("Middle Name", form.middle_name),
        ("Age", form.age),
        ("Gender", form.get_gender_display() if form.gender else ""),
        ("Civil Status", form.get_civil_status_display() if form.civil_status else ""),
        ("Address", form.address),
        ("Date of Birth", form.date_of_birth),
        ("Place of Birth", form.place_of_birth),
        ("Email", form.email_address),
        ("Contact No.", form.contact_number),
        ("Telephone No.", form.telephone_number),
        ("Designation", form.get_designation_display() if form.designation else ""),
        ("Dept / College / Office", form.department_college_office),
        ("Date of Examination", form.date_of_examination),
    ], cols=3)

    _add_data_grid(doc, [
        ("Guardian / Emergency Contact", form.guardian_name),
        ("Contact Number", form.guardian_contact),
    ], cols=2)

    # ── FDI Dental Chart ──
    _add_section_header(doc, "Dental Chart (FDI Notation)")

    teeth = {t.tooth_number: t for t in form.dental_chart.all()}

    # Upper row: 18-11 | 21-28
    upper_right = list(range(18, 10, -1))  # 18,17,...,11
    upper_left = list(range(21, 29))       # 21,22,...,28
    upper_nums = upper_right + upper_left

    # Lower row: 48-41 | 31-38
    lower_right = list(range(48, 40, -1))  # 48,47,...,41
    lower_left = list(range(31, 39))       # 31,32,...,38
    lower_nums = lower_right + lower_left

    for row_label, nums in [("Upper", upper_nums), ("Lower", lower_nums)]:
        table = doc.add_table(rows=2, cols=len(nums))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, num in enumerate(nums):
            _cell_text(table.cell(0, i), str(num), bold=True, size=Pt(7),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
            t = teeth.get(num)
            condition = t.get_condition_display() if t else ""
            cond_short = condition[:3].upper() if condition else ""
            _cell_text(table.cell(1, i), cond_short, size=Pt(6),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_table_borders(table)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

    # ── Soft Tissue Exam ──
    _add_section_header(doc, "Initial Soft Tissue Exam")
    _add_data_grid(doc, [
        ("Lips", form.soft_tissue_lips),
        ("Floor of Mouth", form.soft_tissue_floor_of_mouth),
        ("Palate", form.soft_tissue_palate),
        ("Tongue", form.soft_tissue_tongue),
        ("Neck Nodes", form.soft_tissue_neck_nodes),
    ], cols=3)

    # ── Oral Health Condition ──
    _add_section_header(doc, "Oral Health Condition")
    _add_data_grid(doc, [
        ("Age Last Birthday", form.oral_health_age_last_birthday),
        ("Debris Present", form.presence_of_debris),
        ("Gingival Inflammation", form.inflammation_of_gingiva),
        ("Calculus Present", form.presence_of_calculus),
        ("Orthodontic Treatment", form.under_orthodontic_treatment),
        ("Dentofacial Anomaly", form.dentofacial_anomaly),
    ], cols=3)

    # ── Tooth Count ──
    _add_section_header(doc, "Tooth Count")
    _add_data_grid(doc, [
        ("Teeth Present", form.teeth_present),
        ("Caries-Free", form.caries_free_teeth),
        ("Decayed", form.decayed_teeth),
        ("Missing", form.missing_teeth),
        ("Filled", form.filled_teeth),
        ("Total DMF", form.total_dmf_teeth),
    ], cols=3)

    # ── Periodontal Exam ──
    _add_section_header(doc, "Initial Periodontal Exam")
    _add_data_grid(doc, [
        ("Gingival Inflammation", form.get_gingival_inflammation_display() if form.gingival_inflammation else ""),
        ("Soft Plaque Buildup", form.get_soft_plaque_buildup_display() if form.soft_plaque_buildup else ""),
        ("Hard Calculus Buildup", form.get_hard_calc_buildup_display() if form.hard_calc_buildup else ""),
        ("Stains", form.get_stains_display() if form.stains else ""),
        ("Home Care Effectiveness", form.get_home_care_effectiveness_display() if form.home_care_effectiveness else ""),
        ("Periodontal Condition", form.get_periodontal_condition_display() if form.periodontal_condition else ""),
        ("Periodontal Diagnosis", form.get_periodontal_diagnosis_display() if form.periodontal_diagnosis else ""),
        ("Periodontitis", form.get_periodontitis_display() if form.periodontitis else ""),
        ("Mucogingival Defects", form.mucogingival_defects),
    ], cols=3)

    # ── Clinical Data ──
    _add_section_header(doc, "Clinical Data")
    _add_data_grid(doc, [
        ("Occlusion", form.get_occlusion_display() if form.occlusion else ""),
        ("TMJ Pain", form.tmj_pain),
        ("TMJ Popping", form.tmj_popping),
        ("TMJ Deviation", form.tmj_deviation),
        ("TMJ Tooth Wear", form.tmj_tooth_wear),
    ], cols=3)

    # ── Conditions & Recommendations ──
    _add_section_header(doc, "Conditions & Recommendations")
    conditions = [
        ("Caries Free", form.cond_caries_free),
        ("Poor Oral Hygiene", form.cond_poor_oral_hygiene),
        ("Indicated Restoration", form.cond_indicated_restoration),
        ("Indicated Extraction", form.cond_indicated_extraction),
        ("Gingival Inflammation", form.cond_gingival_inflammation),
        ("Needs Oral Prophylaxis", form.cond_needs_oral_prophylaxis),
        ("Needs Prosthesis", form.cond_needs_prosthesis),
        ("For Endodontic", form.cond_for_endodontic),
        ("For Orthodontic", form.cond_for_orthodontic),
        ("For Sealant", form.cond_for_sealant),
        ("No Treatment Needed", form.cond_no_treatment_needed),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(["Condition", "Status", "Condition", "Status"]):
        _set_cell_shading(table.cell(0, i), LIGHT_FILL)
        _cell_text(table.cell(0, i), hdr, bold=True, size=Pt(8))
    for idx in range(0, len(conditions), 2):
        row = table.add_row()
        n1, v1 = conditions[idx]
        _cell_text(row.cells[0], n1, size=Pt(8))
        _cell_text(row.cells[1], "✓" if v1 else "—", size=Pt(9),
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if idx + 1 < len(conditions):
            n2, v2 = conditions[idx + 1]
            _cell_text(row.cells[2], n2, size=Pt(8))
            _cell_text(row.cells[3], "✓" if v2 else "—", size=Pt(9),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_table_borders(table)

    if form.cond_others and form.cond_others_detail:
        p = doc.add_paragraph()
        _run(p, "Other: ", bold=True, size=Pt(8))
        _run(p, _val(form.cond_others_detail), size=Pt(8))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Remarks
    if form.remarks:
        _add_section_header(doc, "Remarks")
        p = doc.add_paragraph()
        _run(p, _val(form.remarks), size=Pt(9))

    _add_signature_block(
        doc,
        name=_val(form.dentist_name),
        title_text="Dentist",
        license_label="License No.",
        license_no=_val(form.dentist_license_no),
    )

    _add_footer(doc, "F-HSS-20-0003")
    return doc


# ──────────────────────────────────────────────────────────────────
#  Dental Form 2 — Dental Services Request
# ──────────────────────────────────────────────────────────────────

def generate_dental_services(form):
    """Dental Form 2 — Dental Services Request."""
    doc = Document()
    _setup_page(doc)
    _add_letterhead(doc, "Dental Services Request", "Dental Form 2")

    # Personal info
    _add_section_header(doc, "Personal Information")
    _add_data_grid(doc, [
        ("Last Name", form.last_name),
        ("First Name", form.first_name),
        ("Middle Name", form.middle_name),
        ("Address", form.address),
        ("Age", form.age),
        ("Gender", form.get_gender_display() if form.gender else ""),
        ("Date of Birth", form.date_of_birth),
        ("Contact No.", form.contact_number),
        ("Department", form.department),
    ], cols=3)

    # ── Services Checklist ──
    categories = [
        ("Periodontics", [
            ("Oral Prophylaxis", form.perio_oral_prophylaxis, ""),
            ("Scaling & Root Planning", form.perio_scaling_root_planning, ""),
        ]),
        ("Operative Dentistry", [
            ("Class I Restoration", form.oper_class_i, form.oper_class_i_detail),
            ("Class II Restoration", form.oper_class_ii, form.oper_class_ii_detail),
            ("Class III Restoration", form.oper_class_iii, form.oper_class_iii_detail),
            ("Class IV Restoration", form.oper_class_iv, form.oper_class_iv_detail),
            ("Class V Restoration", form.oper_class_v, form.oper_class_v_detail),
            ("Class VI Restoration", form.oper_class_vi, form.oper_class_vi_detail),
            ("Onlay / Inlay", form.oper_onlay_inlay, form.oper_onlay_inlay_detail),
        ]),
        ("Surgery", [
            ("Tooth Extraction", form.surg_tooth_extraction, form.surg_tooth_extraction_detail),
            ("Odontectomy", form.surg_odontectomy, form.surg_odontectomy_detail),
            ("Operculectomy", form.surg_operculectomy, form.surg_operculectomy_detail),
            ("Other Pathological", form.surg_other_pathological, form.surg_other_pathological_detail),
        ]),
        ("Prosthodontics", [
            ("Complete Denture", form.prosth_complete_denture, ""),
            ("RPD", form.prosth_rpd, form.prosth_rpd_detail),
            ("FPD", form.prosth_fpd, form.prosth_fpd_detail),
            ("Single Crown", form.prosth_single_crown, form.prosth_single_crown_detail),
            ("Veneers / Laminates", form.prosth_veneers_laminates, form.prosth_veneers_laminates_detail),
        ]),
        ("Endodontics", [
            ("Anterior", form.endo_anterior, form.endo_anterior_detail),
            ("Posterior", form.endo_posterior, form.endo_posterior_detail),
        ]),
        ("Pediatric Dentistry", [
            ("Fluoride", form.pedo_fluoride, ""),
            ("Sealant", form.pedo_sealant, form.pedo_sealant_detail),
            ("Pulpotomy", form.pedo_pulpotomy, form.pedo_pulpotomy_detail),
            ("SSC", form.pedo_ssc, form.pedo_ssc_detail),
            ("Space Maintainer", form.pedo_space_maintainer, form.pedo_space_maintainer_detail),
        ]),
    ]

    for cat_name, items in categories:
        _add_section_header(doc, cat_name)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hdr in enumerate(["Service", "Selected", "Details"]):
            _set_cell_shading(table.cell(0, i), LIGHT_FILL)
            _cell_text(table.cell(0, i), hdr, bold=True, size=Pt(8))
        for svc_name, checked, detail in items:
            row = table.add_row()
            _cell_text(row.cells[0], svc_name, size=Pt(8))
            _cell_text(row.cells[1], "✓" if checked else "—", size=Pt(9),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row.cells[2], _val(detail, "—"), size=Pt(8))
        _set_table_borders(table)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Treatment status
    if form.currently_undergoing_treatment:
        p = doc.add_paragraph()
        _run(p, "Currently undergoing treatment: ", bold=True, size=Pt(9))
        _run(p, _val(form.currently_undergoing_treatment_detail), size=Pt(9))

    _add_signature_block(
        doc,
        name=_val(form.dentist_name),
        title_text="Dentist",
        license_label="License No.",
        license_no=_val(form.dentist_license_no),
    )

    _add_footer(doc, "Dental Form 2")
    return doc


# ──────────────────────────────────────────────────────────────────
#  F-HSS-20-0004 — Prescription
# ──────────────────────────────────────────────────────────────────

def generate_prescription(rx):
    """F-HSS-20-0004 — Prescription.

    Opens the official HSS-Form0004-Prescription.docx template and fills in
    patient data, prescription items, and physician details.

    The template has two identical tables (side-by-side copies on landscape
    A4).  Each table contains a single cell with 9 paragraphs:
      P0-P3  Letterhead (left untouched)
      P4     Patient fields: NAME / AGE / GENDER / DATE / ADDRESS (blanks)
      P5     Empty — we insert the Rx body + medication items here
      P6     Physician signature line + "MD"
      P7     License No.
      P8     PTR No.
    """
    template_path = TEMPLATES_DIR / "HSS-Form0004-Prescription.docx"
    doc = Document(str(template_path))

    # Build replacement values
    patient_name = _val(rx.patient_name)
    age = _val(rx.age, "")
    gender = rx.get_gender_display() if rx.gender else ""
    date_str = str(rx.date) if rx.date else ""
    address = _val(rx.address, "")
    physician = _val(rx.physician_name, "")
    license_no = _val(rx.license_no, "")
    ptr_no = _val(rx.ptr_no, "")

    # Build Rx body text
    rx_lines = []
    if rx.prescription_body:
        rx_lines.append(rx.prescription_body)
    items = rx.items.all()
    for idx, item in enumerate(items, 1):
        parts = [f"{idx}. {item.medication_name}"]
        if item.dosage:
            parts.append(item.dosage)
        if item.frequency:
            parts.append(f"— {item.frequency}")
        if item.duration:
            parts.append(f"x {item.duration}")
        if item.quantity:
            parts.append(f"#{item.quantity}")
        rx_lines.append(" ".join(parts))
        if item.instructions:
            rx_lines.append(f"   Sig: {item.instructions}")
    rx_body = "\n".join(rx_lines)

    # Fill both copies (Table[0] and Table[1])
    for table in doc.tables:
        cell = table.cell(0, 0)
        paras = cell.paragraphs

        # P[4] — patient info line
        _replace_paragraph_text(
            paras[4],
            f"NAME: {patient_name}    AGE: {age}    GENDER: {gender}    "
            f"DATE: {date_str}    ADDRESS: {address}",
            bold=True,
        )

        # P[5] — Rx body (was empty)
        _replace_paragraph_text(paras[5], rx_body, bold=False)

        # P[6] — physician name + MD
        _replace_paragraph_text(
            paras[6],
            f"{physician}, MD" if physician else "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿MD",
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

        # P[7] — License No.
        _replace_paragraph_text(
            paras[7],
            f"License No. {license_no}" if license_no else "License No. ____________________",
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

        # P[8] — PTR No.
        _replace_paragraph_text(
            paras[8],
            f"PTR No. {ptr_no}" if ptr_no else "PTR No. ______________________",
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

    return doc


def _replace_paragraph_text(paragraph, new_text, bold=None, alignment=None):
    """Replace all runs in a paragraph with a single run containing *new_text*.

    Preserves the font name and size from the first original run so the
    replacement blends seamlessly with the template.
    """
    # Grab formatting from first run before clearing
    font_name = FONT_NAME
    font_size = Pt(10)
    if paragraph.runs:
        first = paragraph.runs[0]
        if first.font.name:
            font_name = first.font.name
        if first.font.size:
            font_size = first.font.size

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""
    # Remove leftover runs from the XML
    p_elem = paragraph._p
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)

    # Add single new run
    run = paragraph.add_run(new_text)
    run.font.name = font_name
    run.font.size = font_size
    if bold is not None:
        run.bold = bold

    if alignment is not None:
        paragraph.alignment = alignment


# ═══════════════════════════════════════════════════════════════════
#  Public convenience — stream as HttpResponse
# ═══════════════════════════════════════════════════════════════════

def doc_to_response(doc, filename):
    """Save a python-docx Document into a Django HttpResponse for download."""
    from django.http import HttpResponse

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response
